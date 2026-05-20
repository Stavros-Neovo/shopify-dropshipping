"""
Erzeugt den Shopify Setup Guide als .docx.
Einmalig laufen lassen: python3 build_setup_guide.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def h(doc, level, text):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
    return p


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return para


def num(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return para


def warning(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("⚠  " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x40, 0x00)
    run.font.size = Pt(11)
    return para


def tip(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("💡  " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x10, 0x60, 0xA0)
    run.font.size = Pt(11)
    return para


doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Seite ")
    add_page_number(fp)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- TITLE ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Shopify Dropshipping Shop")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x10, 0x40, 0x80)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Schritt-für-Schritt Setup-Guide für deinen automatisierten Elektronik/Gadgets-Shop")
sub_run.italic = True
sub_run.font.size = Pt(13)

doc.add_paragraph()

# --- INHALT / EINLEITUNG ---
h(doc, 1, "Einleitung")
p(doc, "Dieser Guide führt dich von einem leeren Shopify-Account bis zu einem voll funktionsfähigen, automatisiertem Dropshipping-Shop. Die zugehörige Automatisierungs-Software wurde bereits gebaut – siehe die Python-Dateien in deinem Workspace-Ordner. Dieser Guide kümmert sich um die Shopify-Seite.")
p(doc, "Geplante Bearbeitungszeit: 4–6 Stunden für die Erst-Einrichtung. Danach läuft der Shop automatisch.")

warning(doc, "WICHTIG – RECHTSHINWEIS: Dropshipping in Deutschland erfordert ein angemeldetes Gewerbe, gültige Pflicht-Pages (Impressum, AGB, Widerrufsbelehrung, Datenschutz) und korrekte MwSt-Behandlung. Dieser Guide gibt praktische Anleitung, ist aber keine Rechtsberatung. Bei größeren Umsätzen unbedingt Steuerberater + ggf. Anwalt einbinden.")

doc.add_paragraph()

# --- KAPITEL 1: SHOP-GRUNDLAGEN ---
h(doc, 1, "1. Shop-Grundeinstellungen")

h(doc, 2, "1.1 Shop-Name & Domain")
num(doc, "Im Shopify-Admin → Settings → General")
num(doc, "Shop-Name eintragen (z.B. \"GadgetVibe\", \"TechFlow\", \"PixelPort\")")
num(doc, "Branche: \"Electronics\" auswählen")
num(doc, "Adresse: deine angemeldete Gewerbe-Adresse")
num(doc, "Zeitzone: (UTC+01:00) Berlin")
num(doc, "Währung: EUR")
tip(doc, "Domain: Anfangs reicht <name>.myshopify.com. Eigene Domain (z.B. .de oder .shop) später dazubuchen – kostet ca. 10–15 €/Jahr.")

h(doc, 2, "1.2 Steuern (KRITISCH für Deutschland)")
num(doc, "Settings → Taxes and duties → European Union")
num(doc, "Tax registration für Deutschland hinzufügen, deine USt-ID eintragen")
num(doc, "Standard-Rate: 19 % (wird automatisch aus Steuersitz ermittelt)")
num(doc, "WICHTIG: \"All prices include tax\" aktivieren – in Deutschland verpflichtend für B2C")
warning(doc, "Wenn du Kleinunternehmer bist (§ 19 UStG), keine USt ausweisen! Dann \"Charge tax on this product\" pro Produkt deaktivieren und einen entsprechenden Hinweis ins Impressum.")

h(doc, 2, "1.3 Zahlungsanbieter aktivieren")
bullet(doc, "Shopify Payments (empfohlen, sofern verfügbar): niedrigste Gebühren, Kreditkarte + Apple Pay + Google Pay")
bullet(doc, "PayPal Express Checkout (Pflicht – über 50 % der DE-Käufer wollen PayPal)")
bullet(doc, "Klarna (Rechnung + Ratenkauf) – essentiell für Conversion in DE")
bullet(doc, "SOFORT / Bank-Überweisung – Klarna-Stack reicht meist")
tip(doc, "Konversion-Faustregel DE: ohne PayPal verlierst du ~30 % der Kunden, ohne Klarna nochmal ~15 %.")

h(doc, 2, "1.4 Versand")
num(doc, "Settings → Shipping and delivery → Manage rates")
num(doc, "General profile → Edit")
num(doc, "Zone \"Deutschland\" anlegen")
num(doc, "Rate hinzufügen: Name \"Standardversand\", Price 4,99 €, Conditions: Weight 0–2 kg")
num(doc, "Weiteres Rate \"NICHT VERSENDBAR\": Weight ab 2 kg, Price 999 € (verhindert versehentliche Bestellungen schwerer Ware)")
tip(doc, "Wenn du auch nach AT/CH liefern willst, eigene Zone mit höherer Rate (z.B. 9,99 €) anlegen. Bei CH zusätzlich Zoll/Einfuhrumsatzsteuer beachten.")

doc.add_paragraph()

# --- KAPITEL 2: PFLICHTSEITEN ---
h(doc, 1, "2. Pflichtseiten (rechtlich erforderlich)")
warning(doc, "Diese Seiten MÜSSEN von jedem deutschen Online-Shop bereitgestellt werden. Ohne sie drohen Abmahnungen ab dem ersten Tag.")

h(doc, 2, "2.1 Impressum")
p(doc, "Settings → Policies → Legal Notice (oder Pages → Add page).")
p(doc, "Mindestinhalte: vollständiger Name, ladungsfähige Adresse, Telefon + Mail, Handelsregister-Nr. (falls vorhanden), USt-ID, Verantwortlicher i.S.d. § 55 RStV, Streitbeilegung-Link der EU-Kommission.")
tip(doc, "Quick-Win: e-recht24.de oder den Trusted Shops-Generator nutzen. Reine 5-Minuten-Sache.")

h(doc, 2, "2.2 Widerrufsbelehrung + Widerrufsformular")
p(doc, "Settings → Policies → Refund policy.")
p(doc, "Verpflichtend nach § 312g BGB: 14-Tage-Widerrufsrecht mit Musterformular.")
warning(doc, "Bei digitalen Inhalten oder versiegelten Waren existieren Sonderregeln. Falls du z.B. Software-Codes verkaufst, brauchst du eine andere Belehrung.")

h(doc, 2, "2.3 AGB (Allgemeine Geschäftsbedingungen)")
p(doc, "Settings → Policies → Terms of service.")
p(doc, "Regelt: Vertragsschluss, Zahlung, Lieferung, Eigentumsvorbehalt, Gewährleistung, Streitbeilegung. Mit Generator erstellen (siehe Tipp 2.1).")

h(doc, 2, "2.4 Datenschutzerklärung")
p(doc, "Settings → Policies → Privacy policy.")
p(doc, "Verpflichtend nach DSGVO. Muss aufführen: alle eingesetzten Tools (Shopify, Klarna, PayPal, Google Analytics, Meta Pixel, Mailchimp, …), Rechtsgrundlage je Verarbeitung, Drittlandtransfers.")
tip(doc, "Tool-Empfehlung: datenschutz-generator.de von Dr. Schwenke – kostet einmalig ca. 10 € und ist anwaltlich gepflegt.")

h(doc, 2, "2.5 Cookie-Banner")
p(doc, "Shopify-Standard reicht NICHT für DSGVO/TTDSG-Konformität. Apps:")
bullet(doc, "Pandectes GDPR (kostenlos für kleine Shops)")
bullet(doc, "Consentmo / Beeketing CookiePro")
bullet(doc, "Usercentrics (Premium, ca. 30 €/Monat)")

doc.add_paragraph()

# --- KAPITEL 3: THEME & DESIGN ---
h(doc, 1, "3. Theme & Design")

h(doc, 2, "3.1 Theme auswählen")
p(doc, "Themes → Theme library → Visit Theme Store.")
p(doc, "Empfehlung für Elektronik/Gadgets:")
bullet(doc, "Dawn (kostenlos, schnell, minimalistisch) – ideal für den Start")
bullet(doc, "Sense (kostenlos, mehr visuelles Storytelling) – falls du viel Content machst")
bullet(doc, "Refresh (49 USD einmalig) – modern, sehr conversion-optimiert")

h(doc, 2, "3.2 Logo & Favicon")
p(doc, "Customize → Theme settings → Logo. Logo-Format: PNG mit transparentem Hintergrund, 300 px breit reicht.")
tip(doc, "Kostenlose Logo-Generatoren: looka.com, hatchful.shopify.com, namelix.com")

h(doc, 2, "3.3 Farben & Typografie")
p(doc, "Für Elektronik bewährt: dunkler Hauptfarbton (Navy, Dark Grey, Schwarz) + ein Akzent (Cyan, Lime, Orange).")
p(doc, "Schriften: Inter, Poppins oder Roboto – modern und sehr gut lesbar.")

h(doc, 2, "3.4 Pflicht-Sections für die Startseite")
num(doc, "Hero-Banner mit klarem Versprechen + CTA-Button (\"Jetzt entdecken\")")
num(doc, "Trust-Badges: Sichere Zahlung, schneller Versand, 14-Tage-Rückgabe")
num(doc, "Featured Collection (4–8 Top-Produkte)")
num(doc, "Vorteile-Sektion (3 Spalten mit Icons: Versand, Garantie, Support)")
num(doc, "Newsletter-Sign-up mit 10 % Erstkäufer-Rabatt")
num(doc, "Kundenstimmen / Reviews (sobald du welche hast)")

doc.add_paragraph()

# --- KAPITEL 4: API-TOKEN ---
h(doc, 1, "4. Admin-API-Token erzeugen (für die Automatisierung)")
p(doc, "Damit das Python-Sync-Skript Produkte automatisch anlegen kann, brauchst du einen Admin-API-Access-Token. So geht's:")
num(doc, "Apps → Develop apps (rechts oben) – ggf. Developer-Modus aktivieren")
num(doc, "\"Create an app\" → Name z.B. \"Dropshipping Sync\"")
num(doc, "\"Configure Admin API scopes\" → folgende Scopes aktivieren:")
bullet(doc, "read_products / write_products")
bullet(doc, "read_inventory / write_inventory")
bullet(doc, "read_orders (NICHT write_orders – wir lesen nur)")
bullet(doc, "read_locations")
num(doc, "Speichern → \"Install app\"")
num(doc, "Auf der App-Seite jetzt: \"Reveal token once\" – Token kopieren und sofort in deine .env-Datei einfügen (Variable: SHOPIFY_ADMIN_TOKEN). Du siehst diesen Token nur EIN MAL!")
warning(doc, "Den Token wie ein Passwort behandeln. Niemals in einem GitHub-Repo committen! Wenn versehentlich geleakt: in der App-Seite \"Uninstall\" + neu erstellen.")

h(doc, 2, "4.1 Shop-Domain finden")
p(doc, "In der Browser-Adresszeile steht z.B. https://meinshop.myshopify.com/admin – \"meinshop\" ist der Wert, der in deine config.yaml unter shopify.shop_domain eingetragen werden muss.")

doc.add_paragraph()

# --- KAPITEL 5: WEBHOOKS ---
h(doc, 1, "5. Auto-Bestellung via Webhook einrichten")

p(doc, "Damit jede bezahlte Bestellung automatisch eine Mail an deinen Lieferanten auslöst, richtest du in Shopify einen Webhook ein, der auf den \"order_forwarder.py\"-Server zeigt.")

h(doc, 2, "5.1 Webhook-Server hosten")
p(doc, "Da der Webhook über das öffentliche Internet erreichbar sein muss, brauchst du einen Server:")
bullet(doc, "Variante A – Render.com (kostenlos für kleine Workloads): GitHub-Repo verbinden, Service-Typ \"Web Service\", Start-Command: python order_forwarder.py serve --port $PORT")
bullet(doc, "Variante B – Railway.app (5 USD/Monat Free-Tier): ähnliche Schritte")
bullet(doc, "Variante C – kleiner VPS (Hetzner Cloud CX11, 4 €/Monat) + systemd-Service")
bullet(doc, "Variante D – ngrok für TEMPORÄRE Tests (lokal, nicht produktionsfest)")

h(doc, 2, "5.2 Webhook in Shopify einrichten")
num(doc, "Settings → Notifications → Webhooks")
num(doc, "\"Create webhook\" → Event: \"Order payment\" (also wenn die Bestellung BEZAHLT ist – nicht bei \"Order creation\", sonst kommen auch unbezahlte Bestellungen durch)")
num(doc, "Format: JSON")
num(doc, "URL: https://deine-server-domain/webhook")
num(doc, "API-Version: 2024-10")
num(doc, "Webhook-Secret kopieren und in .env als SHOPIFY_WEBHOOK_SECRET speichern")

h(doc, 2, "5.3 Test-Bestellung")
num(doc, "Lege eine Testbestellung mit \"Bogus Gateway\" auf (Settings → Payments → Manage → Test Mode)")
num(doc, "Schau in den Server-Logs, ob der Webhook ankommt")
num(doc, "In outbox/ sollte eine .eml-Datei liegen (auto_send: false in config.yaml)")
num(doc, "Inhalt prüfen → wenn ok: auto_send: true setzen")

doc.add_paragraph()

# --- KAPITEL 6: GO-LIVE CHECKLISTE ---
h(doc, 1, "6. Go-Live Checkliste")

h(doc, 2, "6.1 Letzte Tests vor dem Schalter umlegen")
num(doc, "python sync.py --dry-run – Log prüfen, mind. 20 Produkte sollten korrekt importiert werden")
num(doc, "python sync.py --live – Erst-Import durchführen (ca. 250 Produkte, ~10 Minuten)")
num(doc, "Shopify Admin → Products: Sind alle Produkte da? Bilder geladen? Preise sinnvoll?")
num(doc, "Test-Bestellung mit Bogus Gateway durchführen")
num(doc, "Outbox-Mail prüfen: Adresse korrekt? SKU richtig? Versand-Anrede sauber?")
num(doc, "Webhook auf \"Order paid\" testet eine echte Live-Karte für ~1 €, dann sofort stornieren")

h(doc, 2, "6.2 Marketing-Vorbereitung")
bullet(doc, "Google Merchant Center: Feed in Shopify einrichten (App \"Google & YouTube\")")
bullet(doc, "Meta Business Suite verknüpfen für Facebook/Instagram Ads")
bullet(doc, "Mindestens 5–10 Produktfotos eigenständig erstellen (z.B. für Reels/TikTok)")
bullet(doc, "Newsletter-Welcome-Flow in Shopify Email oder Klaviyo")

h(doc, 2, "6.3 Erste 30 Tage")
bullet(doc, "Täglich Sync-Logs prüfen (logs/sync_*.log)")
bullet(doc, "Wöchentlich: Filter-Gründe analysieren – sind sinnvoll viele Produkte aus? Margen ok?")
bullet(doc, "Bei der ersten echten Bestellung: Lieferanten-Antwort abwarten, Versand-Tracking checken")
bullet(doc, "Reviews einsammeln: Loox oder Judge.me App – schickt automatisch Review-Anfragen 14 Tage nach Lieferung")

doc.add_paragraph()

# --- KAPITEL 7: TROUBLESHOOTING ---
h(doc, 1, "7. Troubleshooting")

h(doc, 2, "Sync schlägt fehl mit 401 Unauthorized")
p(doc, "→ Token in .env stimmt nicht oder hat keine ausreichenden Scopes. App in Shopify uninstallen + neu erstellen.")

h(doc, 2, "Produkte werden ohne Bilder importiert")
p(doc, "→ Bild-URLs in der CSV prüfen. Shopify lehnt Bilder ab, die a) hinter Login liegen, b) > 20 MB sind, c) keine gültige Endung (.jpg/.png/.webp) haben.")

h(doc, 2, "VK-Preis ist zu niedrig / zu hoch")
p(doc, "→ python pricing.py ausführen, sehen wie die Staffel rechnet. Dann tiers in config.yaml anpassen.")

h(doc, 2, "Bestellung kommt rein, aber Lieferant erhält keine Mail")
p(doc, "→ Erste Anlaufstelle: outbox/-Ordner prüfen. Wenn dort eine .eml liegt, ist auto_send: false in config.yaml gesetzt. Wenn nichts da, Webhook-URL in Shopify prüfen + Server-Logs.")

h(doc, 2, "GitHub Action timeoutet")
p(doc, "→ max_products_per_run in config.yaml senken, z.B. von 250 auf 100. Mehrere Läufe pro Stunde sind besser als 1 langer Lauf.")

doc.add_paragraph()

# --- ABSCHLUSS ---
h(doc, 1, "8. Nächste Schritte nach dem Go-Live")
bullet(doc, "Monat 1: Daten sammeln. Welche Produkte konvertieren? Welche Quellen liefern Traffic?")
bullet(doc, "Monat 2: Top-Performer mit Meta-Ads / Google-Ads skalieren (Budget: 10–20 €/Tag pro Produkt)")
bullet(doc, "Monat 3: Sortiment kuratieren – schlechte Produkte aus filtern (allowed_categories oder excluded_keywords)")
bullet(doc, "Monat 4+: E-Mail-Marketing-Flows ausbauen (Abandoned Cart, Post-Purchase, Win-Back)")
bullet(doc, "Sobald > 10 Bestellungen/Tag: Steuerberater einbinden, Buchhaltungs-Tool wie sevDesk oder lexoffice via Shopify-Connector")

doc.add_paragraph()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
final_run = final.add_run("Viel Erfolg mit deinem Shop! 🚀")
final_run.bold = True
final_run.font.size = Pt(13)

# Save
output = "Shopify_Setup_Guide.docx"
doc.save(output)
print(f"Erstellt: {output}")
